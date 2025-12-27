import requests
from bs4 import BeautifulSoup
import openai
import random
import os
import time
from datetime import datetime
from docx import Document
from docx.shared import Inches, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# OpenAI API 설정
openai.api_key = "api_key"

# 매일경제 메인 페이지 URL
main_url = 'https://www.mk.co.kr/news/economy/'

# 1. 인기 기사 링크 가져오기
def get_popular_articles():
    response = requests.get(main_url)
    soup = BeautifulSoup(response.text, 'html.parser')

    # 기사 링크 추출
    article_links = []
    for i in range(1, 11):
        article = soup.select(f'#list_area > li:nth-child({i}) > a')
        if article:
            article_links.append(article[0]['href'])
    return article_links

# 2. 기사 제목과 본문 추출
def get_article_details(url):
    response = requests.get(url)
    soup = BeautifulSoup(response.text, 'html.parser')

    # 제목 추출
    title_tag = soup.select_one('#container > section.contents > div.news_detail_head_group.type_none_bg > section > div > div > div > h2')
    title = title_tag.get_text().strip() if title_tag else '제목을 찾을 수 없습니다.'

    # 본문 추출
    content_tag = soup.select_one('div.news_cnt_detail_wrap')
    content = content_tag.get_text() if content_tag else '본문을 찾을 수 없습니다.'
    
    return title, content

# 3. OpenAI로 요약 요청
def summarize_article_with_openai(content):
    try:
        response = openai.chat.completions.create(
            model="gpt-3.5-turbo",  # GPT 모델 사용
            messages=[{
                "role": "user",
                "content": f"다음 기사의 본문을 보고 그 기사에 대한 내 생각을 만들어서 의견을 3~4줄로 요약해주세요:\n\n{content}"
            }]
        )
        ai_response = response.choices[0].message.content
        return ai_response
    except Exception as e:
        return f"Error: {str(e)}"

# 메인 함수
def main():
    # 인기 기사 링크 10개 가져오기
    article_links = get_popular_articles()

    # 3개 랜덤 기사 선택
    selected_links = random.sample(article_links, 3)

    # 기존 파일 경로
    file_path = r"C:\Users\Desktop\경제신문.docx"
    
    # 기존 파일이 있으면 열고, 없으면 새로 생성
    if os.path.exists(file_path):
        try:
            document = Document(file_path)
            print(f"기존 파일을 열었습니다: {file_path}")
            # 기존 내용 뒤에 구분선 추가
            document.add_paragraph("=" * 50)
            document.add_paragraph()
        except Exception as e:
            print(f"기존 파일을 열 수 없습니다. 새 파일을 생성합니다: {e}")
            document = Document()
            document.add_heading('경제 기사 요약', 0)
    else:
        document = Document()
        document.add_heading('경제 기사 요약', 0)
        print(f"새 파일을 생성합니다: {file_path}")

    # 날짜 현재 시간
    current_date = datetime.now().strftime("%Y-%m-%d")

    # 각 기사마다 별도의 2열 표 생성 (헤더 | 내용)
    for idx, url in enumerate(selected_links, 1):
        print(f"기사 {idx} URL: {url}")
        
        # 기사 제목과 본문 가져오기
        title, content = get_article_details(url)

        # 본문을 요약
        summary = summarize_article_with_openai(content)

        print(f"기사 {idx} 제목: {title}")
        print(f"기사 {idx} 요약: {summary}")

        # 2열 표 생성 (헤더 | 내용)
        table = document.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        # 표 레이아웃을 고정(fixed)으로 설정
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # 기존 tblLayout 제거
        for elem in tblPr:
            if elem.tag == qn('w:tblLayout'):
                tblPr.remove(elem)
                break
        
        # 고정 레이아웃 설정
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)
        
        # 열 너비 설정: 왼쪽 헤더 열은 최소(2cm), 오른쪽 내용 열은 나머지(14cm)
        header_width = Cm(2.0)   # 헤더 열 - 글자수만큼만
        content_width = Cm(14.0)  # 내용 열 - 나머지 공간
        
        # 각 행의 셀 너비 설정
        for row in table.rows:
            row.cells[0].width = header_width   # 왼쪽 열 (헤더)
            row.cells[1].width = content_width  # 오른쪽 열 (내용)
        
        # 표에 데이터 추가
        # 첫 번째 행: 날짜
        table.rows[0].cells[0].text = '날짜'
        table.rows[0].cells[1].text = current_date
        
        # 두 번째 행: 기사 제목
        table.rows[1].cells[0].text = '기사 제목'
        table.rows[1].cells[1].text = title
        
        # 세 번째 행: 기사 링크
        table.rows[2].cells[0].text = '기사 링크'
        table.rows[2].cells[1].text = url
        
        # 네 번째 행: 기사 요약
        table.rows[3].cells[0].text = '기사 요약'
        table.rows[3].cells[1].text = summary
        
        # 기사 사이에 빈 줄 추가
        if idx < len(selected_links):
            document.add_paragraph()

    # 결과를 Word 문서로 저장 (에러 처리 포함)
    saved = False
    max_attempts = 5
    
    for attempt in range(max_attempts):
        try:
            document.save(file_path)
            print(f"\n✓ 성공: '{os.path.basename(file_path)}' 파일에 내용이 추가되었습니다.")
            print(f"  저장 위치: {file_path}")
            saved = True
            break
            
        except PermissionError as e:
            if attempt < max_attempts - 1:
                print(f"권한 오류 발생. 재시도 중... ({attempt+1}/{max_attempts-1})")
                time.sleep(1)  # 1초 대기 후 재시도
            else:
                print(f"\n✗ 오류: 파일을 저장할 수 없습니다. (에러 13: Permission denied)")
                print("\n다음 사항을 확인해주세요:")
                print("1. Word에서 '경제신문.docx' 파일이 열려있는지 확인하고 닫아주세요.")
                print("2. 파일 탐색기에서 해당 파일이 열려있는지 확인하고 닫아주세요.")
                print("3. 바이러스 백신 프로그램이 파일을 차단하고 있는지 확인해주세요.")
                print("4. 파일에 대한 쓰기 권한이 있는지 확인해주세요.")
        except Exception as e:
            print(f"\n✗ 오류 발생: {e}")
            if attempt < max_attempts - 1:
                time.sleep(1)
            else:
                break
    
    if not saved:
        print("\n파일 저장에 실패했습니다. 프로그램을 종료합니다.")

if __name__ == '__main__':
    main()
